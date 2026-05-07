"""
Creates a professional A4 Word document template for Oades Consultancy Ltd.
Logo: PNG screenshot embedded directly on a matching #1A2B3E navy banner.
"""
import io, os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PNG   = r"C:\Users\timke\OneDrive\Pictures\Screenshots 1\Screenshot 2026-04-30 215118.png"
OUTPUT     = r"C:\Users\timke\oades\oades_consultancy_template.docx"
TMP_HEADER = r"C:\Users\timke\oades\_header_banner.png"
TMP_FOOTER = r"C:\Users\timke\oades\_footer_banner.png"

# Brand colours (sampled from the logo PNG)
BG_COLOR  = (26, 43, 62)          # #1A2B3E – dark navy
GOLD_COLOR = (184, 151, 61)       # #B8973D
WHITE_COLOR = (255, 255, 255)

# A4 at 300 DPI: 2480 × 3508 px. Header/footer banner width = 2480 px.
DPI         = 300
A4_W_PX     = 2480   # full A4 width at 300 DPI
HDR_H_PX    = 280    # header band height in pixels
FTR_H_PX    = 160    # footer band height


# ---------------------------------------------------------------------------
# Banner image builders
# ---------------------------------------------------------------------------

def prepare_logo(logo_path):
    """
    Convert a white-background dark-ink logo to a white-ink transparent-background image
    suitable for compositing onto the dark navy banner.
    """
    import numpy as np
    img  = Image.open(logo_path).convert("RGBA")
    arr  = np.array(img, dtype=np.float32)

    # Lightness of each pixel (0 = black ink, 255 = white background)
    lightness = arr[:, :, :3].mean(axis=2)

    # Build alpha: fully transparent where near-white (background), opaque where dark (ink)
    alpha = np.clip((255 - lightness) * (255 / 200), 0, 255).astype(np.uint8)

    # Recolour all ink to pure white
    out       = np.zeros_like(arr, dtype=np.uint8)
    out[:, :, 0] = 255
    out[:, :, 1] = 255
    out[:, :, 2] = 255
    out[:, :, 3] = alpha

    return Image.fromarray(out, "RGBA")


def make_header_banner(logo_path, out_path):
    """Create the full-width header banner PNG."""
    canvas = Image.new("RGB", (A4_W_PX, HDR_H_PX), BG_COLOR)
    draw   = ImageDraw.Draw(canvas)

    # Process logo: strip white bg, recolour ink to white
    logo = prepare_logo(logo_path)

    # Scale logo to fit left ~45 % of banner height, vertically centred
    # (this logo is portrait so constrain by height)
    max_logo_h = HDR_H_PX - 20
    max_logo_w = int(A4_W_PX * 0.45)
    ratio      = min(max_logo_w / logo.width, max_logo_h / logo.height)
    new_w      = int(logo.width  * ratio)
    new_h      = int(logo.height * ratio)
    logo       = logo.resize((new_w, new_h), Image.LANCZOS)

    logo_x = 60
    logo_y = (HDR_H_PX - new_h) // 2
    canvas.paste(logo, (logo_x, logo_y), logo)

    # "German Legal Consultancy" tagline – right-aligned, vertically centred
    tagline = "German Legal Consultancy"
    # Try to load a serif/italic font; fall back to default
    font = None
    font_candidates = [
        r"C:\Windows\Fonts\calibrii.ttf",   # Calibri Italic
        r"C:\Windows\Fonts\calibrib.ttf",   # Calibri Bold
        r"C:\Windows\Fonts\calibri.ttf",    # Calibri Regular
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for fc in font_candidates:
        if os.path.exists(fc):
            try:
                font = ImageFont.truetype(fc, size=62)
                break
            except Exception:
                pass
    if font is None:
        font = ImageFont.load_default()

    bbox   = draw.textbbox((0, 0), tagline, font=font)
    txt_w  = bbox[2] - bbox[0]
    txt_h  = bbox[3] - bbox[1]
    txt_x  = A4_W_PX - txt_w - 60
    txt_y  = (HDR_H_PX - txt_h) // 2

    # Subtle shadow
    draw.text((txt_x + 2, txt_y + 2), tagline, font=font, fill=(10, 20, 30))
    # Gold text
    draw.text((txt_x, txt_y), tagline, font=font, fill=GOLD_COLOR)

    # Gold bottom rule (3 px)
    rule_y = HDR_H_PX - 8
    draw.rectangle([0, rule_y, A4_W_PX, HDR_H_PX], fill=GOLD_COLOR)

    canvas.save(out_path, "PNG", dpi=(DPI, DPI))
    print(f"  Header banner saved ({canvas.size[0]}×{canvas.size[1]} px)")


def make_footer_banner(out_path):
    """Create the full-width footer banner PNG with contact info."""
    canvas = Image.new("RGB", (A4_W_PX, FTR_H_PX), BG_COLOR)
    draw   = ImageDraw.Draw(canvas)

    # Gold top rule
    draw.rectangle([0, 0, A4_W_PX, 8], fill=GOLD_COLOR)

    # Fonts
    font_bold = font_reg = None
    for path, size in [
        (r"C:\Windows\Fonts\calibrib.ttf", 46),
        (r"C:\Windows\Fonts\arialbd.ttf",  46),
    ]:
        if os.path.exists(path):
            try:
                font_bold = ImageFont.truetype(path, size=size)
                break
            except Exception:
                pass
    for path, size in [
        (r"C:\Windows\Fonts\calibri.ttf",  38),
        (r"C:\Windows\Fonts\arial.ttf",    38),
    ]:
        if os.path.exists(path):
            try:
                font_reg = ImageFont.truetype(path, size=size)
                break
            except Exception:
                pass
    if font_bold is None:
        font_bold = ImageFont.load_default()
    if font_reg is None:
        font_reg = ImageFont.load_default()

    company = "Oades Consultancy Ltd."
    details = ("+44 (0) 730 9555 935   |   2 St. Thomas Square, Salisbury, SP1 1BA, "
               "United Kingdom   |   www.oadesconsultancy.co.uk")

    # Company name centred
    cb = draw.textbbox((0, 0), company, font=font_bold)
    cw = cb[2] - cb[0]
    draw.text(((A4_W_PX - cw) // 2, 28), company, font=font_bold, fill=WHITE_COLOR)

    # Detail line centred
    db = draw.textbbox((0, 0), details, font=font_reg)
    dw = db[2] - db[0]
    draw.text(((A4_W_PX - dw) // 2, 88), details, font=font_reg, fill=(180, 195, 210))

    canvas.save(out_path, "PNG", dpi=(DPI, DPI))
    print(f"  Footer banner saved ({canvas.size[0]}×{canvas.size[1]} px)")


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def clear_header_footer(hf):
    for p in list(hf.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(hf.tables):
        t._element.getparent().remove(t._element)


# A4 exact Word twips (Word standard values)
A4_TWIPS   = 11906   # page width
MARGIN_TWP = 1440    # 1 inch = 1440 twips (2.54 cm margin)


def build_banner_table(container, banner_path, page_w_in):
    """
    Add a truly full-width, zero-padding banner image table to a header/footer.
    Uses exact A4 twip values and strips every source of internal whitespace.
    """
    table = container.add_table(rows=1, cols=1, width=Inches(page_w_in))
    tbl   = table._tbl

    # ── Table properties ───────────────────────────────────────────────────
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Exact page width
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(A4_TWIPS))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Negative left indent to reach left page edge
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"),    str(-MARGIN_TWP))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    # No cell spacing
    tblCellSp = OxmlElement("w:tblCellSpacing")
    tblCellSp.set(qn("w:w"),    "0")
    tblCellSp.set(qn("w:type"), "dxa")
    tblPr.append(tblCellSp)

    # Zero default cell margins (overrides Word's built-in padding)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    "0")
        el.set(qn("w:type"), "dxa")
        tblCellMar.append(el)
    tblPr.append(tblCellMar)

    # No table borders
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tblBdr.append(el)
    tblPr.append(tblBdr)

    # ── Cell properties ────────────────────────────────────────────────────
    cell = table.cell(0, 0)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Exact cell width
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(A4_TWIPS))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

    # No cell borders
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tcBdr.append(el)
    tcPr.append(tcBdr)

    # Zero cell internal margins
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    "0")
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

    # Cell background
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1A2B3E")
    tcPr.append(shd)

    # ── Paragraph inside cell ──────────────────────────────────────────────
    para = cell.paragraphs[0]
    pPr  = para._p.get_or_add_pPr()

    # Zero paragraph spacing
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"),   "0")
    pSp.set(qn("w:after"),    "0")
    pSp.set(qn("w:line"),     "240")
    pSp.set(qn("w:lineRule"), "auto")
    pPr.append(pSp)

    # Zero paragraph indents
    pInd = OxmlElement("w:ind")
    pInd.set(qn("w:left"),  "0")
    pInd.set(qn("w:right"), "0")
    pPr.append(pInd)

    # Paragraph background
    pShd = OxmlElement("w:shd")
    pShd.set(qn("w:val"),   "clear")
    pShd.set(qn("w:color"), "auto")
    pShd.set(qn("w:fill"),  "1A2B3E")
    pPr.append(pShd)

    # Picture — exactly page width
    para.add_run().add_picture(banner_path, width=Inches(page_w_in))

    return table


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def build():
    MARGIN_CM  = 2.54
    PAGE_W_IN  = 8.27          # A4
    BODY_W_IN  = PAGE_W_IN - (MARGIN_CM / 2.54) * 2   # 6.27 in

    print("Building header banner …")
    make_header_banner(LOGO_PNG, TMP_HEADER)
    print("Building footer banner …")
    make_footer_banner(TMP_FOOTER)

    doc = Document()
    section = doc.sections[0]
    section.page_height        = Cm(29.7)
    section.page_width         = Cm(21.0)
    section.left_margin        = Cm(MARGIN_CM)
    section.right_margin       = Cm(MARGIN_CM)
    section.top_margin         = Cm(3.0)
    section.bottom_margin      = Cm(2.5)
    section.header_distance    = Cm(0)
    section.footer_distance    = Cm(0)

    # "Different first page" so the banner header only appears on page 1
    section.different_first_page_header_footer = True

    # ── HEADER: first page only ──────────────────────────────────────────────
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    clear_header_footer(first_header)
    build_banner_table(first_header, TMP_HEADER, PAGE_W_IN)

    # Pages 2+ get an empty header
    other_header = section.header
    other_header.is_linked_to_previous = False
    clear_header_footer(other_header)
    other_header.add_paragraph()

    # ── FOOTER: every page (first page + default) ────────────────────────────
    def attach_footer(footer_obj):
        footer_obj.is_linked_to_previous = False
        clear_header_footer(footer_obj)
        build_banner_table(footer_obj, TMP_FOOTER, PAGE_W_IN)

    attach_footer(section.first_page_footer)
    attach_footer(section.footer)

    # ── BODY STYLES ─────────────────────────────────────────────────────────
    NAVY      = RGBColor(0x1A, 0x2B, 0x3E)
    GOLD      = RGBColor(0xB8, 0x97, 0x3D)
    DARK_GREY = RGBColor(0x3A, 0x3A, 0x3A)

    doc.styles["Normal"].font.name       = "Calibri"
    doc.styles["Normal"].font.size       = Pt(11)
    doc.styles["Normal"].font.color.rgb  = DARK_GREY

    h1 = doc.styles["Heading 1"]
    h1.font.name      = "Georgia"
    h1.font.size      = Pt(16)
    h1.font.color.rgb = NAVY
    h1.font.bold      = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    pPr1 = h1.element.get_or_add_pPr()
    pBdr1 = OxmlElement("w:pBdr")
    bot1  = OxmlElement("w:bottom")
    bot1.set(qn("w:val"),   "single")
    bot1.set(qn("w:sz"),    "12")
    bot1.set(qn("w:space"), "4")
    bot1.set(qn("w:color"), "B8973D")
    pBdr1.append(bot1)
    pPr1.append(pBdr1)

    h2 = doc.styles["Heading 2"]
    h2.font.name      = "Georgia"
    h2.font.size      = Pt(13)
    h2.font.color.rgb = NAVY
    h2.font.bold      = True

    # ── PLACEHOLDER BODY ────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after  = Pt(4)
    tr = title_p.add_run("[Document Title]")
    tr.font.name      = "Georgia"
    tr.font.size      = Pt(22)
    tr.font.bold      = True
    tr.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(30)
    sr = sub_p.add_run("[Date]  ·  [Reference]  ·  CONFIDENTIAL")
    sr.font.name      = "Calibri"
    sr.font.size      = Pt(10)
    sr.font.italic    = True
    sr.font.color.rgb = GOLD

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Begin your document content here. This template uses Oades Consultancy Ltd. "
        "brand colours — deep navy (#1A2B3E) and gold — for a professional legal document style."
    ).paragraph_format.space_after = Pt(8)

    doc.add_heading("2. Section Heading", level=1)
    doc.add_paragraph(
        "Continue with your content. Additional headings and body text follow "
        "the same typographic hierarchy."
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

    # Clean up temp files
    for f in (TMP_HEADER, TMP_FOOTER):
        try:
            os.remove(f)
        except Exception:
            pass


if __name__ == "__main__":
    build()
